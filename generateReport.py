import constants
import os
import shutil
import tkinter
import openai
import time
from os import path
from docx import Document
from tkinter import filedialog
from openai import OpenAI

filePath = lambda file: path.abspath(file.name)

def check_word_file():
    return path.isfile( f"./{constants.REPORT_FILE}" )

def clone_word_file( src, output ):
    if path.isfile( output ):
        os.remove( output )
    return shutil.copy( src, output )


def edit_word_file(file, param, value):
    document = Document( file )
    match param:
        case "author":
            table = document.tables[0]
            author_row = table.rows[1]
            author_cell = author_row.cells[1]
            for paragraph in author_cell.paragraphs:
                if paragraph.text == f"<{param}>":
                    paragraph.text = value 
            pass
        case "task_number":
            table = document.tables[0]
            author_row = table.rows[0]
            author_cell = author_row.cells[3]
            for paragraph in author_cell.paragraphs:
                if paragraph.text == f"<{param}>":
                    paragraph.text = value 
            pass
        case "task_title":
            table = document.tables[0]
            author_row = table.rows[2]
            author_cell = author_row.cells[1]
            for paragraph in author_cell.paragraphs:
                if paragraph.text == f"<{param}>":
                    paragraph.text = value
            pass
        case "group_number":
            table = document.tables[0]
            group_row = table.rows[1]
            group_cell = group_row.cells[3]
            for paragraph in group_cell.paragraphs:
                if paragraph.text == f"<{param}>":
                    paragraph.text = value
            pass
        case _:
            for paragraph in document.paragraphs:
                if paragraph.text == f"<{param}>":
                    paragraph.text = value
            pass
    document.save(file)

def start_report( userData ):
    startTime = time.time()
    openaiClient = OpenAI(
        api_key=userData.openAISecret
    )
    
    if not check_word_file():
        print( f"No report file. File should be in program folder with name: {constants.REPORT_FILE }" )
        return
    
    window = tkinter.Tk()
    window.withdraw()
    window.title( "Wybierz plik zadania (kodu)" )
    codeFile = filedialog.askopenfile(mode='r', filetypes=[("C/C++ files", ".c .cc .cpp .h ")])
    
    if not codeFile:
        print( "Please select valid code file next time." )
        return
    print( codeFile.name )
    
    taskNumber = input( "Podaj numer zadania: " )
    taskTitle = input("Podaj nazwe zadania: " )

    outputFile = clone_word_file( f"./{constants.REPORT_FILE}", outputName := constants.OUTPUT_FILE( taskNumber ) )

    edit_word_file( outputName, "author", f"{userData.name} {userData.surname}" )
    edit_word_file( outputName, "group_number", userData.groupNumber )
    edit_word_file( outputName, "task_number", taskNumber )
    edit_word_file( outputName, "task_title", taskTitle )

    realisationMethod = input( "Podaj metodę realizacji: ")
    edit_word_file( outputName, "realisation_method", realisationMethod )

    entryData = input( "Podaj dane wejsciowe zadania: ")
    edit_word_file( outputName, "entry_data", entryData )

    exitData = input( "Podaj dane wyjściowe zadania: ")
    edit_word_file( outputName, "exit_data", exitData )

    codeContent = open(filePath(codeFile), encoding="UTF-8").readlines()
    edit_word_file( outputName, "source_code", codeContent )
    print( codeContent )
    
    """response = openaiClient.chat.completions.create(
        model="gpt-3.5-turbo-instruct",
        messages=f"{codeContent} \n lista kroków"
    )
    print(response)"""
                                        
    summary = input( "Podaj wnioski dotyczące programu:" ) 
    edit_word_file( outputName, "summary", summary )

    print( f"Finished generating report. Took: {time.time()-startTime}s")

    return